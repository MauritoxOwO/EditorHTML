from __future__ import annotations

import argparse
import html
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
PAGE_MARGIN_CM = 2.54
TABLE_WIDTH_CM = 15.4
COL_WIDTHS_CM = [2.1, 3.2, 6.0, 4.1]
FONT_FAMILY = "Calibri"
FONT_SIZE_PT = 11
HEADER_ROW_HEIGHT_PT = 28
BODY_ROW_HEIGHT_PT = 35


def set_cell_margins(cell, top: int = 57, start: int = 85, bottom: int = 57, end: int = 85) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_cm: float) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(round(width_cm / 2.54 * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(width_cm / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def format_paragraph(paragraph, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    for run in paragraph.runs:
        run.bold = bold
        run.font.name = FONT_FAMILY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FAMILY)
        run.font.size = Pt(FONT_SIZE_PT)


def build_docx(output_path: Path, rows: int) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(PAGE_MARGIN_CM)
    section.right_margin = Cm(PAGE_MARGIN_CM)
    section.bottom_margin = Cm(PAGE_MARGIN_CM)
    section.left_margin = Cm(PAGE_MARGIN_CM)
    section.start_type = WD_SECTION.NEW_PAGE

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    set_table_width(table, TABLE_WIDTH_CM)

    headers = ["ID", "Estado", "Descripcion", "Responsable"]
    table.rows[0].height = Pt(HEADER_ROW_HEIGHT_PT)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = header_cells[idx]
        set_cell_width(cell, COL_WIDTHS_CM[idx])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cell.text = text
        format_paragraph(cell.paragraphs[0], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_index in range(1, rows + 1):
        row = table.add_row()
        row.height = Pt(BODY_ROW_HEIGHT_PT)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row_cells = row.cells
        values = [
            f"{row_index:03d}",
            "Activo" if row_index % 3 else "Revision",
            (
                "Linea de tabla Word para validar paginacion, cortes por fila y altura "
                f"estable del editor. Registro {row_index} con texto suficiente para envolver."
            ),
            f"Equipo {((row_index - 1) % 5) + 1}",
        ]
        for idx, value in enumerate(values):
            cell = row_cells[idx]
            set_cell_width(cell, COL_WIDTHS_CM[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = value
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(cell.paragraphs[0], align=align)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_word_html(output_path: Path, rows: int) -> None:
    colgroup = "".join(f'<col style="width:{width:.2f}cm">' for width in COL_WIDTHS_CM)
    header = "".join(
        f'<td style="width:{COL_WIDTHS_CM[idx]:.2f}cm;border:solid windowtext 1.0pt;'
        "padding:2.85pt 4.25pt 2.85pt 4.25pt;vertical-align:top\">"
        f'<p class="MsoNormal" align="center" style="margin:0cm;text-align:center;'
        f'line-height:normal"><b><span style="font-family:{FONT_FAMILY};font-size:{FONT_SIZE_PT}.0pt">'
        f"{html.escape(label)}</span></b></p></td>"
        for idx, label in enumerate(["ID", "Estado", "Descripcion", "Responsable"])
    )
    body_rows = []
    for row_index in range(1, rows + 1):
        values = [
            f"{row_index:03d}",
            "Activo" if row_index % 3 else "Revision",
            (
                "Linea de tabla Word para validar paginacion, cortes por fila y altura "
                f"estable del editor. Registro {row_index} con texto suficiente para envolver."
            ),
            f"Equipo {((row_index - 1) % 5) + 1}",
        ]
        cells = []
        for idx, value in enumerate(values):
            align = "center" if idx in (0, 1, 3) else "left"
            cells.append(
                f'<td style="width:{COL_WIDTHS_CM[idx]:.2f}cm;border:solid windowtext 1.0pt;'
                'padding:2.85pt 4.25pt 2.85pt 4.25pt;vertical-align:top">'
                f'<p class="MsoNormal" align="{align}" style="margin:0cm;'
                f'text-align:{align};line-height:normal">'
                f'<span style="font-family:{FONT_FAMILY};font-size:{FONT_SIZE_PT}.0pt">'
                f"{html.escape(value)}</span></p></td>"
            )
        body_rows.append(f'<tr style="height:{BODY_ROW_HEIGHT_PT}.0pt">' + "".join(cells) + "</tr>")

    word_html = f"""<!DOCTYPE html>
<html xmlns:o="urn:schemas-microsoft-com:office:office"
      xmlns:w="urn:schemas-microsoft-com:office:word"
      xmlns="http://www.w3.org/TR/REC-html40">
<head>
<meta charset="utf-8">
<style>
@page WordSection1 {{
  size: 210mm 297mm;
  margin: 25.4mm 25.4mm 25.4mm 25.4mm;
}}
div.WordSection1 {{ page: WordSection1; }}
p.MsoNormal {{
  margin: 0cm;
  font-size: {FONT_SIZE_PT}.0pt;
  font-family: {FONT_FAMILY};
}}
table.MsoTableGrid {{
  border-collapse: collapse;
  border: none;
  margin-left: auto;
  margin-right: auto;
}}
</style>
</head>
<body>
<!--StartFragment-->
<div class="WordSection1">
<table class="MsoTableGrid" align="center" border="1" cellspacing="0" cellpadding="0"
       style="border-collapse:collapse;border:none;margin-left:auto;margin-right:auto;width:{TABLE_WIDTH_CM:.2f}cm">
<colgroup>{colgroup}</colgroup>
<tbody>
<tr style="height:{HEADER_ROW_HEIGHT_PT}.0pt">{header}</tr>
{''.join(body_rows)}
</tbody>
</table>
</div>
<!--EndFragment-->
</body>
</html>
"""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(word_html, encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--rows", type=int, default=76)
    parser.add_argument("--docx", type=Path, default=Path("fixtures/word-table-4-pages.docx"))
    parser.add_argument(
        "--html", type=Path, default=Path("web/public/fixtures/word-table-4-pages.html")
    )
    args = parser.parse_args()

    build_docx(args.docx, args.rows)
    build_word_html(args.html, args.rows)
    print(f"Generated {args.docx} and {args.html} with {args.rows} body rows.")


if __name__ == "__main__":
    main()
